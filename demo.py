#!/usr/bin/env python
import sys
import json
import numpy as np
from docx import Document
import matplotlib.pyplot as plt
from risk_score_chartjs import generate_risk_score_chartjs
import random

path = sys.argv[1]


def read_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()


# Function to calculate the sum of field_a.value and field_b.value


def calculate_sum(input_json):
    field_a_value = input_json["field_a"]["value"]
    field_b_value = input_json["field_b"]["value"]

    if field_b_value == 11 and field_b_value == 11:
        raise Exception("Test error triggered by passing 11:11")
    sum_value = field_a_value + field_b_value
    return sum_value

def generateLabels():
    return []

def generateData(val):
    return [val+random.uniform(-0.5, 0.5) for x in range(12)]

COLORS = {
  'red': 'rgb(255, 99, 132)',
  'orange': 'rgb(255, 159, 64)',
  'yellow': 'rgb(255, 205, 86)',
  'green': 'rgb(75, 192, 192)',
  'blue': 'rgb(54, 162, 235)',
  'purple': 'rgb(153, 102, 255)',
  'grey': 'rgb(201, 203, 207)'
}

BACK_COLORS = {
  'red': 'rgba(255, 99, 132, 0.1)',
  'orange': 'rgba(255, 159, 64, 0.1)',
  'yellow': 'rgba(255, 205, 86, 0.1)',
  'green': 'rgba(75, 192, 192, 0.1)',
  'blue': 'rgba(54, 162, 235, 0.1)',
  'purple': 'rgba(153, 102, 255, 0.1)',
  'grey': 'rgba(201, 203, 207, 0.1)'
}

BAR_COLORS = {
  'red': 'rgba(255, 99, 132, 0.5)',
  'orange': 'rgba(255, 159, 64, 0.5)',
  'yellow': 'rgba(255, 205, 86, 0.5)',
  'green': 'rgba(75, 192, 192, 0.5)',
  'blue': 'rgba(54, 162, 235, 0.5)',
  'purple': 'rgba(153, 102, 255, 0.5)',
  'grey': 'rgba(201, 203, 207, 0.5)'
}
def generate_polar_chart():
    data = {
      'labels': [f'risks_{x}' for x in range(5)],
      'datasets': [
        {
          'label': 'Dataset 1',
          'data': [random.uniform(1.0, 100.0) for _ in range(5)],
          'backgroundColor': [
            BAR_COLORS['red'],
            BAR_COLORS['orange'],
            BAR_COLORS['yellow'],
            BAR_COLORS['green'],
            BAR_COLORS['blue']
          ]
        }
      ]
    }
    config = {
      'type': 'polarArea',
      'data': data,
      'options': {
        'responsive': 'true',
        'plugins': {
          'legend': {
            'position': 'top',
          },
          'title': {
            'display': 'true',
            'text': 'Chart.js Polar Area Chart'
          }
        }
      },
    }
    return config

def generate_radar_chart():
    data = {
      'labels': [f'L{x}' for x in range(12)],
      'datasets': [
        {
          'label': 'D0',
          'data': generateData(1),
          'borderColor': COLORS['red'],
          'backgroundColor': BACK_COLORS['red']
        },
        {
          'label': 'D1',
          'data': generateData(2),
          'borderColor': COLORS['orange'],
          'hidden': 'true',
          'backgroundColor': BACK_COLORS['orange'],
          'fill': '-1'
        },
        {
          'label': 'D2',
          'data': generateData(3),
          'borderColor': COLORS['yellow'],
          'backgroundColor': BACK_COLORS['yellow'],
          'fill': '1'
        },
        {
          'label': 'D3',
          'data': generateData(4),
          'borderColor': COLORS['green'],
          'backgroundColor': BACK_COLORS['green'],
          'fill': 'false'
        },
        {
          'label': 'D4',
          'data': generateData(5),
          'borderColor': COLORS['blue'],
          'backgroundColor': BACK_COLORS['blue'],
          'fill': '-1'
        },
        {
          'label': 'D5',
          'data': generateData(6),
          'borderColor': COLORS['purple'],
          'backgroundColor': BACK_COLORS['purple'],
          'fill': '-1'
        },
        {
          'label': 'D6',
          'data': generateData(7),
          'borderColor': COLORS['grey'],
          'backgroundColor': BACK_COLORS['grey'],
          'fill': {'value': '85'}
        }
      ]
    }
    config = {
      'type': 'radar',
      'data': data,
      'options': {
        'plugins': {
          'filler': {
            'propagate': 'false'
          },
          'samples-filler-analyser': {
            'target': 'chart-analyser'
          }
        },
        'interaction': {
          'intersect': 'false'
        }
      }
    }
    return config

def generate_line_chart_data(num_points, lang):
    num_points = int(num_points)
    if num_points < 0:
        num_points = 1

    x_values = np.linspace(0, 100, num_points)
    y_values1 = np.sin(x_values)
    y_values2 = np.cos(x_values)

    if lang == 'lv':
        label = 'Līnijas grafiks no HPC'
    else:
        label = 'Line chart from HPC'

    line_chart_data = {
        'type': 'line',
        'data': {
            'labels': list(x_values),
            'datasets': [
                {
                    'label': 'Sin',
                    'data': list(y_values1),
                },
                {
                    'label': 'Cos',
                    'data': list(y_values2),
                }
            ]
        },
        'options': {
            "plugins": {
                "legend": {
                    "position": 'top',
                },
                "title": {
                    "display": "true",
                    "text": label
                }
            },
            "scales": {
                "x": {
                    "type": "linear",
                    "position": "bottom"
                },
                "y": {
                    "type": "linear"
                }
            }
        }
    }
    return line_chart_data


def generate_test_msword(input_json, value, lang):
    text = read_file(path + input_json["field_user_text"]["filename"])

    a = input_json["field_a"]["value"]
    b = input_json["field_b"]["value"]

    doc = Document()
    # Add a paragraph with the sample text

    if lang == "lv":
        doc.add_paragraph("Mēs veicāc summas aprēķiņu: " + str(a) + "+" + str(b) + "=" + str(value))
        doc.add_paragraph("Jūs iesniedzāt šādu teksta dokumentu: ")
        doc.add_paragraph(text)
        doc.save(path + "/output/files/my_word_lv.docx")
        return "/output/files/my_word_lv.docx"
    else:
        doc.add_paragraph("We calculated a sum: " + str(a) + "+" + str(b) + "=" + str(value))
        doc.add_paragraph("You have provided the text document below: ")
        doc.add_paragraph(text)
        doc.add_paragraph("UAT tests dokumenta satura modifikācijai")
        doc.save(path + "/output/files/my_word_en.docx")
        return "/output/files/my_word_en.docx"


def generate_test_dt(input_json, lang):
    if lang == 'lv':
        oper_label = "Operācija"
        res_label = 'Rezultāts'
        caption = "Aprēķini"
    else:
        oper_label = "Operation"
        res_label = "Result"
        caption = "Calculations"

    data = [
            [input_json["field_a"]["value"], '+', input_json["field_b"]["value"],
             input_json["field_a"]["value"] + input_json["field_b"]["value"]],
            [input_json["field_a"]["value"], '-', input_json["field_b"]["value"],
             input_json["field_a"]["value"] - input_json["field_b"]["value"]],
            [input_json["field_a"]["value"], '*', input_json["field_b"]["value"],
             input_json["field_a"]["value"] * input_json["field_b"]["value"]],
            [input_json["field_a"]["value"], '/', input_json["field_b"]["value"],
             round(input_json["field_a"]["value"] / input_json["field_b"]["value"] * 100) / 100],
        ]
    for i in range(100):
        data.append([str(i), str(i+1), str(i+2), str(i+3)])

    dt = {
        'caption': caption,
        'columns': [
            {'title': 'A'},
            {'title': oper_label},
            {'title': 'B'},
            {'title': res_label}
        ],
        'data': data
    }
    return dt


def generate_heatmap(locale):
    plt.imshow(np.random.rand(10, 20))
    outname = f"/output/files/heatmap_{locale}.png"
    plt.savefig(f"{path}{outname}")
    plt.close()
    return outname

def generate_test_chart(input_json, sum, lang):
    a = input_json["field_a"]["value"]
    b = input_json["field_b"]["value"]

    # Data for the bars
    categories = ['a', 'b', 'y=a+b']
    values = [a, b, sum]

    # Create a bar chart
    a = np.linspace(0, 100)
    b = np.linspace(100, 200)
    plt.scatter(a, b)

    if lang == 'lv':
        # Add title and labels
        plt.title('Joslu diagrammas piemērs')
        plt.xlabel('Mainīgie')
        plt.ylabel('Vērtības')
    else:
        # Add title and labels
        lang = 'en'
        plt.title('A bar chart example')
        plt.xlabel('Variables')
        plt.ylabel('Values')

    # Save the chart as a PNG file
    plt.savefig(path + "/output/files/" + lang + "_chart.png")
    return "/output/files/" + lang + "_chart.png"


# Read input JSON file
with open(path + "/input/input.json", "r") as f:
    input_data = json.load(f)

# Calculate sum
sum_value = calculate_sum(input_data)

# Create output JSON object
output_data = {
    "values": {
        "sum": sum_value
    },
    "datatables": {
        "tableLv": generate_test_dt(input_data, 'lv'),
        "tableEn": generate_test_dt(input_data, 'en')
    },
    "chartjs": {
        "myLine_lv": generate_line_chart_data(sum_value, 'lv'),
        "myLine_en": generate_line_chart_data(sum_value, 'en'),
        "risk_lv": generate_risk_score_chartjs(score=input_data["field_a"]["value"], lang='lv'),
        "risk_en": generate_risk_score_chartjs(score=input_data["field_a"]["value"], lang='en'),
        "radar": generate_radar_chart(),
        "polar": generate_polar_chart(),
    },
    "files": {
        "MyWord_lv": generate_test_msword(input_data, sum_value, 'lv'),
        "MyWord_en": generate_test_msword(input_data, sum_value, 'en'),
        "MyChart_lv": generate_test_chart(input_data, sum_value, 'lv'),
        "MyChart_en": generate_test_chart(input_data, sum_value, 'en'),
        "heatmap_en": generate_heatmap('en'),
        "heatmap_lv": generate_heatmap('lv'),
    }
}

# Write output JSON to STDOUT
print(json.dumps(output_data))
